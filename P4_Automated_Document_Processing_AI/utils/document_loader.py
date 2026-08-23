from pathlib import Path
from pypdf import PdfReader
from docx import Document

def load_document(uploaded_file):
    ext = Path(uploaded_file.name).suffix.lower()
    if ext == ".pdf":
        text = "\n\n".join((p.extract_text() or "") for p in PdfReader(uploaded_file).pages)
    elif ext == ".docx":
        doc = Document(uploaded_file)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        text = "\n".join(parts)
    elif ext == ".txt":
        text = uploaded_file.read().decode("utf-8", errors="replace")
    else:
        raise ValueError("Only PDF, DOCX and TXT files are supported.")
    if not text.strip():
        raise ValueError("No readable text was found in the document.")
    return text.strip()
